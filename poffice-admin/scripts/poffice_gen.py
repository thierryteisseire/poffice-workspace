import os
from docx import Document
from openpyxl import Workbook, load_workbook
from reportlab.pdfgen import canvas
from reportlab.lib.pagesizes import letter

class PofficeDocumentGenerator:
    @staticmethod
    def create_word(path, title, content):
        doc = Document()
        doc.add_heading(title, 0)
        doc.add_paragraph(content)
        doc.save(path)
        return path

    @staticmethod
    def update_word(path, new_paragraph):
        doc = Document(path)
        doc.add_paragraph(new_paragraph)
        doc.save(path)
        return path

    @staticmethod
    def create_excel(path, sheet_name, data):
        """data should be a list of lists: [['Col1', 'Col2'], [1, 2]]"""
        wb = Workbook()
        ws = wb.active
        ws.title = sheet_name
        for row in data:
            ws.append(row)
        wb.save(path)
        return path

    @staticmethod
    def update_excel(path, data):
        wb = load_workbook(path)
        ws = wb.active
        for row in data:
            ws.append(row)
        wb.save(path)
        return path

    @staticmethod
    def create_pdf(path, title, content):
        c = canvas.Canvas(path, pagesize=letter)
        width, height = letter
        c.setFont("Helvetica-Bold", 16)
        c.drawString(72, height - 72, title)
        c.setFont("Helvetica", 12)
        
        # Simple line splitter for content
        textobject = c.beginText(72, height - 100)
        for line in content.split('\n'):
            textobject.textLine(line)
        c.drawText(textobject)
        
        c.showPage()
        c.save()
        return path

if __name__ == "__main__":
    import sys
    gen = PofficeDocumentGenerator()
    if len(sys.argv) > 1:
        action = sys.argv[1]
        if action == "create-word":
            # python3 poffice_gen.py create-word test.docx "My Title" "My Content"
            print(gen.create_word(sys.argv[2], sys.argv[3], sys.argv[4]))
        elif action == "create-pdf":
            print(gen.create_pdf(sys.argv[2], sys.argv[3], sys.argv[4]))
